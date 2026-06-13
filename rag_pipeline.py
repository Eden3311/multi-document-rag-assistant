"""
rag_pipeline.py
---------------
Core RAG pipeline using FAISS vector store + Groq LLM.

PERMANENT FIX for "Cannot send a request, as the client has been closed":
The HuggingFace embedding model is now loaded ONCE and cached as a
module-level singleton (via functools.lru_cache). Previously it was
recreated on every single Streamlit rerun, which spawned a new internal
HTTP client each time. After enough reruns, a stale/closed client was
reused internally and crashed with this exact error.

We also force HuggingFace Hub into OFFLINE mode after the model is cached
locally for the first time — this stops ALL background network/version
checks, making the app fully usable without internet (except the Groq
LLM call, which needs internet).
"""

import os
import gc
import json
import shutil
from functools import lru_cache
from pathlib import Path
from typing import List, Tuple

# ── Document loaders ─────────────────────────────────────────────────────────
import fitz
import docx

# ── LangChain ────────────────────────────────────────────────────────────────
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq

# Ollama import is optional — only needed for offline mode.
# Wrapped in try/except so the app doesn't crash if langchain-ollama
# isn't installed (user can still use online/Groq mode).
try:
    from langchain_ollama import ChatOllama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False


# ════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ════════════════════════════════════════════════════════════════════════════

FAISS_STORE_DIR = "./faiss_store"
MANIFEST_PATH   = "./faiss_store/manifest.json"   # tracks all indexed source filenames
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
GROQ_MODEL      = "llama-3.3-70b-versatile"
OLLAMA_MODEL    = "llama3.2"
CHUNK_SIZE      = 800
CHUNK_OVERLAP   = 150
TOP_K           = 8


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT PARSING
# ════════════════════════════════════════════════════════════════════════════

def _parse_pdf(file_path: str) -> str:
    text = ""
    with fitz.open(file_path) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                text += f"\n[Page {page_num}]\n{page_text}"
    return text


def _parse_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _parse_txt(file_path: str) -> str:
    try:
        return Path(file_path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return Path(file_path).read_text(encoding="latin-1")


def parse_document(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: '{ext}'.")


# ════════════════════════════════════════════════════════════════════════════
# CHUNKING
# ════════════════════════════════════════════════════════════════════════════

def chunk_text(raw_text: str, source_name: str) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return [
        Document(
            page_content=chunk,
            metadata={"source": source_name, "chunk_index": i}
        )
        for i, chunk in enumerate(splitter.split_text(raw_text))
    ]


# ════════════════════════════════════════════════════════════════════════════
# EMBEDDINGS — SINGLETON (loaded only ONCE per Python process)
# ════════════════════════════════════════════════════════════════════════════

@lru_cache(maxsize=1)
def _get_embeddings() -> HuggingFaceEmbeddings:
    """
    Load the sentence-transformer embedding model exactly ONCE and
    cache it for the lifetime of the Streamlit process.

    lru_cache(maxsize=1) ensures that no matter how many times this
    function is called across Streamlit reruns, the SAME model object
    (and its internal HTTP client) is reused — never recreated.

    On the FIRST call, the model may need to contact HuggingFace Hub
    to check for updates. After that first successful load, we flip
    HF_HUB_OFFLINE=1 so all future processes skip the network check
    entirely (faster startup + works offline).
    """
    model = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    # After first successful load, force offline mode for any
    # subsequent internal calls in this process (no more HTTP clients
    # being spun up for version checks)
    os.environ["HF_HUB_OFFLINE"] = "1"
    os.environ["TRANSFORMERS_OFFLINE"] = "1"

    return model


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT MANIFEST — ground-truth list of all indexed source files
# ════════════════════════════════════════════════════════════════════════════
#
# similarity_search() only returns chunks RELEVANT to the question — so for
# meta-questions like "how many documents do you have?", retrieval may miss
# a document whose content doesn't closely match the query embedding.
#
# To answer such questions accurately, we keep a simple JSON manifest of
# every unique source filename ever added to the store. This is updated on
# every ingest and wiped on clear — completely independent of similarity
# search relevance.

def _load_manifest() -> List[str]:
    """Return the list of all indexed source filenames (empty if none)."""
    if not Path(MANIFEST_PATH).exists():
        return []
    try:
        with open(MANIFEST_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def _save_manifest(sources: List[str]) -> None:
    """Persist the list of unique source filenames to disk."""
    Path(FAISS_STORE_DIR).mkdir(parents=True, exist_ok=True)
    with open(MANIFEST_PATH, "w", encoding="utf-8") as f:
        json.dump(sources, f, indent=2)


def _add_to_manifest(documents: List[Document]) -> None:
    """Add any new source filenames from `documents` to the manifest."""
    existing = _load_manifest()
    for doc in documents:
        src = doc.metadata.get("source", "Unknown")
        if src not in existing:
            existing.append(src)
    _save_manifest(existing)


def get_all_indexed_sources() -> List[str]:
    """
    Public helper — returns the full, accurate list of all documents
    ever indexed (regardless of retrieval relevance).
    """
    return _load_manifest()


# ════════════════════════════════════════════════════════════════════════════
# FAISS STORE
# ════════════════════════════════════════════════════════════════════════════

def add_documents_to_store(documents: List[Document]) -> None:
    """
    Embed chunks and save to FAISS index on disk.
    Merges with existing index if one exists.
    Also updates the document manifest (ground-truth source list).
    """
    embeddings = _get_embeddings()
    store_path = Path(FAISS_STORE_DIR)

    if store_path.exists():
        try:
            existing  = FAISS.load_local(
                FAISS_STORE_DIR,
                embeddings,
                allow_dangerous_deserialization=True,
            )
            new_store = FAISS.from_documents(documents, embeddings)
            existing.merge_from(new_store)
            existing.save_local(FAISS_STORE_DIR)
            _add_to_manifest(documents)
            return
        except Exception:
            # Corrupt store — wipe and start fresh
            _force_delete(store_path)

    store_path.mkdir(parents=True, exist_ok=True)
    new_store = FAISS.from_documents(documents, embeddings)
    new_store.save_local(FAISS_STORE_DIR)
    _add_to_manifest(documents)


def load_vector_store():
    """
    Load FAISS index from disk.
    Returns None if no documents indexed yet or if loading fails.
    """
    store_path = Path(FAISS_STORE_DIR)
    if not store_path.exists():
        return None
    try:
        embeddings = _get_embeddings()
        return FAISS.load_local(
            FAISS_STORE_DIR,
            embeddings,
            allow_dangerous_deserialization=True,
        )
    except Exception:
        return None


def _force_delete(store_path: Path) -> None:
    """Forcefully delete a folder on Windows even if files are soft-locked."""
    shutil.rmtree(store_path, ignore_errors=True)
    if store_path.exists():
        for f in store_path.iterdir():
            try:
                f.unlink(missing_ok=True)
            except Exception:
                pass
        try:
            store_path.rmdir()
        except Exception:
            pass


def clear_vector_store() -> None:
    """Wipe the entire FAISS store. Safe on Windows."""
    gc.collect()
    _force_delete(Path(FAISS_STORE_DIR))


# ════════════════════════════════════════════════════════════════════════════
# PROMPT
# ════════════════════════════════════════════════════════════════════════════

RAG_PROMPT_TEMPLATE = """You are an expert research assistant. Answer the user's
question using the information from the context below.

IMPORTANT — Document Count:
The user has uploaded EXACTLY {num_docs} document(s) in total:
{doc_list}
If asked "how many documents" or "what documents do you have", answer with
EXACTLY {num_docs} and list ONLY the filenames shown above.
Note: the CONTEXT section below may only contain excerpts from SOME of these
documents (the ones most relevant to the question) — this does not mean the
other documents don't exist.

Rules:
- Give a clear, detailed answer if it is in the context.
- Always state which document the information came from.
- If the context lacks the answer to a CONTENT question say:
  "I couldn't find a clear answer in the uploaded documents."
- Never invent information not present in the context.
- Use bullet points when listing multiple points.

────────────────────────────────────────
CONTEXT (relevant excerpts):
{context}
────────────────────────────────────────

QUESTION: {question}

ANSWER:"""


# ════════════════════════════════════════════════════════════════════════════
# ANSWER QUESTION
# ════════════════════════════════════════════════════════════════════════════

def answer_question(
    question: str,
    mode: str = "online",
    groq_api_key: str = "",
) -> Tuple[str, List[dict]]:
    """
    Full RAG query:
      1. Load FAISS index (cached embeddings, no new clients)
      2. Retrieve top-k relevant chunks
      3. Build prompt with context
      4. Call LLM — either Groq (online) or Ollama (offline)
      5. Return answer + source citations

    Parameters
    ----------
    question     : str   User's question.
    mode         : str   "online"  -> uses Groq cloud API (needs internet + API key)
                          "offline" -> uses local Ollama model (no internet needed)
    groq_api_key : str   Required only when mode == "online".
    """
    vector_store = load_vector_store()
    if vector_store is None:
        return "⚠️ No documents indexed yet. Please upload at least one document first.", []

    source_docs = vector_store.similarity_search(question, k=TOP_K)

    context = "\n\n---\n\n".join(
        f"[Source: {doc.metadata.get('source', 'Unknown')}]\n{doc.page_content}"
        for doc in source_docs
    )

    # Use the MANIFEST for an accurate, complete list of all indexed
    # documents — independent of what similarity_search happened to retrieve
    # for this specific question.
    all_sources  = get_all_indexed_sources()
    doc_list_str = "\n".join(f"- {src}" for src in all_sources)

    prompt = RAG_PROMPT_TEMPLATE.format(
        context=context,
        question=question,
        num_docs=len(all_sources),
        doc_list=doc_list_str,
    )

    # ── Select LLM based on mode ─────────────────────────────────────────
    if mode == "offline":
        if not OLLAMA_AVAILABLE:
            return (
                "⚠️ Offline mode requires the 'langchain-ollama' package.\n"
                "Install it with: pip install langchain-ollama\n"
                "Also make sure Ollama is running and you've pulled a model:\n"
                "ollama pull llama3.2",
                [],
            )
        try:
            llm = ChatOllama(model=OLLAMA_MODEL, temperature=0.2)
            response = llm.invoke(prompt)
        except Exception as e:
            return (
                f"⚠️ Could not reach local Ollama server.\n\n"
                f"Make sure Ollama is installed and running, and that you've run:\n"
                f"ollama pull {OLLAMA_MODEL}\n\n"
                f"Error details: {str(e)}",
                [],
            )
    else:
        # Online mode — Groq
        if not groq_api_key:
            return "⚠️ Please enter your Groq API key in the sidebar for online mode.", []

        llm = ChatGroq(
            model=GROQ_MODEL,
            api_key=groq_api_key,
            temperature=0.2,
            max_tokens=1024,
        )
        response = llm.invoke(prompt)

    answer = response.content

    seen, sources = set(), []
    for doc in source_docs:
        src = doc.metadata.get("source", "Unknown")
        if src not in seen:
            seen.add(src)
            sources.append({
                "source":  src,
                "snippet": doc.page_content[:200].replace("\n", " "),
            })

    return answer, sources


# ════════════════════════════════════════════════════════════════════════════
# INGEST FILE
# ════════════════════════════════════════════════════════════════════════════

def ingest_file(file_path: str, original_name: str = None) -> Tuple[int, str]:
    """
    Parse → chunk → embed → store a single uploaded file.

    Parameters
    ----------
    file_path     : str   Path to the temp file on disk (used for parsing).
    original_name : str   The user's original filename (used for metadata
                           and citations). If not provided, falls back to
                           the temp file's name.
    """
    source_name = original_name if original_name else Path(file_path).name
    raw_text    = parse_document(file_path)

    if not raw_text.strip():
        raise ValueError(f"No text could be extracted from '{source_name}'.")

    documents = chunk_text(raw_text, source_name)
    add_documents_to_store(documents)
    return len(documents), source_name
